from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from pydantic import BaseModel, Field

from app.domain.evidence import DocumentSourceType


class ParsedSegment(BaseModel):
    ordinal: int
    page_number: int | None = None
    text: str
    metadata: dict[str, Any] = Field(default_factory=dict)


class ParsedDocument(BaseModel):
    source_type: DocumentSourceType
    parser_name: str
    segments: list[ParsedSegment]

    @property
    def full_text(self) -> str:
        return "\n".join(segment.text for segment in self.segments if segment.text).strip()


def parse_document(filename: str, raw_bytes: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()

    if suffix in {".txt", ".md"}:
        text = raw_bytes.decode("utf-8")
        return ParsedDocument(
            source_type=DocumentSourceType.TEXT,
            parser_name="plain_text",
            segments=[ParsedSegment(ordinal=0, text=text)],
        )

    if suffix == ".pdf":
        pdf = fitz.open(stream=raw_bytes, filetype="pdf")
        try:
            segments = [
                ParsedSegment(
                    ordinal=index,
                    page_number=index + 1,
                    text=page.get_text("text").strip(),
                )
                for index, page in enumerate(pdf)
            ]
        finally:
            pdf.close()
        return ParsedDocument(
            source_type=DocumentSourceType.PDF,
            parser_name="pymupdf",
            segments=segments,
        )

    if suffix == ".docx":
        document = Document(BytesIO(raw_bytes))
        segments = [
            ParsedSegment(ordinal=index, text=paragraph.text.strip())
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip()
        ]
        return ParsedDocument(
            source_type=DocumentSourceType.DOCX,
            parser_name="python-docx",
            segments=segments,
        )

    if suffix in {".png", ".jpg", ".jpeg"}:
        return ParsedDocument(
            source_type=DocumentSourceType.IMAGE,
            parser_name="multimodal_required",
            segments=[],
        )

    return ParsedDocument(
        source_type=DocumentSourceType.UNKNOWN,
        parser_name="unsupported",
        segments=[],
    )


def extract_text(filename: str, raw_bytes: bytes) -> str:
    return parse_document(filename, raw_bytes).full_text
